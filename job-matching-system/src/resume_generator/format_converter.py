"""
Resume format converter for generating PDF, DOCX, and TXT formats.
"""
import os
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import io

logger = logging.getLogger(__name__)

class FormatConverter:
    def __init__(self, output_directory: str = 'data/generated_resumes'):
        self.output_directory = output_directory
        os.makedirs(output_directory, exist_ok=True)
    
    def convert_to_all_formats(self, resume_data: Dict[str, Any]) -> Dict[str, str]:
        """
        Convert resume to PDF, DOCX, and TXT formats.
        
        Args:
            resume_data: Structured resume data from ResumeBuilder
            
        Returns:
            Dictionary with paths to generated files
        """
        metadata = resume_data.get('metadata', {})
        job_title = metadata.get('job_title', 'Position')
        company = metadata.get('company', 'Company')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Clean filename
        safe_job_title = self._clean_filename(job_title)
        safe_company = self._clean_filename(company)
        base_filename = f"{safe_job_title}_{safe_company}_{timestamp}"
        
        file_paths = {}
        
        try:
            # Generate DOCX
            docx_path = os.path.join(self.output_directory, f"{base_filename}.docx")
            self._create_docx(resume_data, docx_path)
            file_paths['docx'] = docx_path
            
            # Generate PDF
            pdf_path = os.path.join(self.output_directory, f"{base_filename}.pdf")
            self._create_pdf(resume_data, pdf_path)
            file_paths['pdf'] = pdf_path
            
            # Generate TXT
            txt_path = os.path.join(self.output_directory, f"{base_filename}.txt")
            self._create_txt(resume_data, txt_path)
            file_paths['txt'] = txt_path
            
            logger.info(f"Resume generated in all formats for {job_title} at {company}")
            
        except Exception as e:
            logger.error(f"Error converting resume formats: {e}")
            raise
        
        return file_paths
    
    def _clean_filename(self, text: str) -> str:
        """Clean text for safe filename usage."""
        # Remove/replace unsafe characters
        unsafe_chars = '<>:"/\\|?*'
        for char in unsafe_chars:
            text = text.replace(char, '_')
        
        # Limit length and remove extra spaces
        text = ' '.join(text.split())[:50]
        return text
    
    def _create_docx(self, resume_data: Dict[str, Any], file_path: str):
        """Create DOCX format resume."""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header with name and contact info
        header = doc.add_heading('', 0)
        header_run = header.add_run('YOUR NAME')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info (placeholder)
        contact = doc.add_paragraph()
        contact.add_run('Email: your.email@email.com | Phone: (xxx) xxx-xxxx | LinkedIn: linkedin.com/in/yourprofile')
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Professional Summary
        if resume_data.get('professional_summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(resume_data['professional_summary'])
            doc.add_paragraph()
        
        # Core Competencies
        if resume_data.get('core_competencies'):
            doc.add_heading('CORE COMPETENCIES', level=1)
            competencies = resume_data['core_competencies']
            if isinstance(competencies, list):
                # Format as bullet points in columns
                competencies_text = ' • '.join(competencies)
                doc.add_paragraph(competencies_text)
            else:
                doc.add_paragraph(str(competencies))
            doc.add_paragraph()
        
        # Professional Experience
        if resume_data.get('professional_experience'):
            doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
            doc.add_paragraph(resume_data['professional_experience'])
            doc.add_paragraph()
        
        # Education
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=1)
            doc.add_paragraph(resume_data['education'])
            doc.add_paragraph()
        
        # Additional Sections
        if resume_data.get('additional_sections'):
            doc.add_paragraph(resume_data['additional_sections'])
        
        # Save document
        doc.save(file_path)
        logger.info(f"DOCX resume saved to {file_path}")
    
    def _create_pdf(self, resume_data: Dict[str, Any], file_path: str):
        """Create PDF format resume."""
        doc = SimpleDocTemplate(file_path, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        # Build content
        content = []
        
        # Header
        content.append(Paragraph("YOUR NAME", title_style))
        content.append(Paragraph("Email: your.email@email.com | Phone: (xxx) xxx-xxxx | LinkedIn: linkedin.com/in/yourprofile", body_style))
        content.append(Spacer(1, 12))
        
        # Professional Summary
        if resume_data.get('professional_summary'):
            content.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
            content.append(Paragraph(resume_data['professional_summary'], body_style))
        
        # Core Competencies
        if resume_data.get('core_competencies'):
            content.append(Paragraph("CORE COMPETENCIES", heading_style))
            competencies = resume_data['core_competencies']
            if isinstance(competencies, list):
                competencies_text = ' • '.join(competencies)
                content.append(Paragraph(competencies_text, body_style))
            else:
                content.append(Paragraph(str(competencies), body_style))
        
        # Professional Experience
        if resume_data.get('professional_experience'):
            content.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
            # Split experience into paragraphs for better formatting
            experience_lines = resume_data['professional_experience'].split('\n')
            for line in experience_lines:
                if line.strip():
                    content.append(Paragraph(line.strip(), body_style))
        
        # Education
        if resume_data.get('education'):
            content.append(Paragraph("EDUCATION", heading_style))
            content.append(Paragraph(resume_data['education'], body_style))
        
        # Additional Sections
        if resume_data.get('additional_sections'):
            additional_lines = resume_data['additional_sections'].split('\n')
            for line in additional_lines:
                if line.strip():
                    content.append(Paragraph(line.strip(), body_style))
        
        # Build PDF
        doc.build(content)
        logger.info(f"PDF resume saved to {file_path}")
    
    def _create_txt(self, resume_data: Dict[str, Any], file_path: str):
        """Create plain text format resume."""
        lines = []
        
        # Header
        lines.extend([
            "YOUR NAME",
            "Email: your.email@email.com | Phone: (xxx) xxx-xxxx",
            "LinkedIn: linkedin.com/in/yourprofile",
            "",
            "=" * 60,
            ""
        ])
        
        # Professional Summary
        if resume_data.get('professional_summary'):
            lines.extend([
                "PROFESSIONAL SUMMARY",
                "-" * 20,
                resume_data['professional_summary'],
                ""
            ])
        
        # Core Competencies
        if resume_data.get('core_competencies'):
            lines.extend([
                "CORE COMPETENCIES",
                "-" * 17,
            ])
            competencies = resume_data['core_competencies']
            if isinstance(competencies, list):
                for comp in competencies:
                    lines.append(f"• {comp}")
            else:
                lines.append(str(competencies))
            lines.append("")
        
        # Professional Experience
        if resume_data.get('professional_experience'):
            lines.extend([
                "PROFESSIONAL EXPERIENCE",
                "-" * 23,
                resume_data['professional_experience'],
                ""
            ])
        
        # Education
        if resume_data.get('education'):
            lines.extend([
                "EDUCATION",
                "-" * 9,
                resume_data['education'],
                ""
            ])
        
        # Additional Sections
        if resume_data.get('additional_sections'):
            lines.extend([
                "ADDITIONAL INFORMATION",
                "-" * 22,
                resume_data['additional_sections']
            ])
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        logger.info(f"TXT resume saved to {file_path}")
    
    def create_resume_from_template(self, resume_data: Dict[str, Any], template_path: str) -> str:
        """
        Create resume using a DOCX template.
        
        Args:
            resume_data: Resume data structure
            template_path: Path to DOCX template file
            
        Returns:
            Path to generated resume file
        """
        if not os.path.exists(template_path):
            logger.warning(f"Template not found: {template_path}. Using default format.")
            return self.convert_to_all_formats(resume_data)['docx']
        
        try:
            # Load template
            doc = Document(template_path)
            
            # Replace placeholders in template
            replacements = {
                '{{PROFESSIONAL_SUMMARY}}': resume_data.get('professional_summary', ''),
                '{{CORE_COMPETENCIES}}': self._format_competencies_for_template(resume_data.get('core_competencies', [])),
                '{{PROFESSIONAL_EXPERIENCE}}': resume_data.get('professional_experience', ''),
                '{{EDUCATION}}': resume_data.get('education', ''),
                '{{ADDITIONAL_SECTIONS}}': resume_data.get('additional_sections', '')
            }
            
            # Replace text in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
            
            # Replace text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, replacement in replacements.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, replacement)
            
            # Save customized resume
            metadata = resume_data.get('metadata', {})
            job_title = self._clean_filename(metadata.get('job_title', 'Position'))
            company = self._clean_filename(metadata.get('company', 'Company'))
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            output_path = os.path.join(self.output_directory, f"{job_title}_{company}_{timestamp}_template.docx")
            doc.save(output_path)
            
            logger.info(f"Template-based resume saved to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating resume from template: {e}")
            # Fallback to default format
            return self.convert_to_all_formats(resume_data)['docx']
    
    def _format_competencies_for_template(self, competencies) -> str:
        """Format competencies for template insertion."""
        if isinstance(competencies, list):
            return ' • '.join(competencies)
        return str(competencies)
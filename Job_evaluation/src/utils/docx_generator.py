import logging
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO

class DocxGenerator:
    """Utility class for generating DOCX documents"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def create_feedback_document(self, job: Dict, evaluation: Dict) -> BytesIO:
        """Create a feedback document in DOCX format"""
        try:
            doc = Document()
            
            # Document title
            title = doc.add_heading('Job Evaluation Feedback', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Job information section
            doc.add_heading('Job Information', level=1)
            
            job_table = doc.add_table(rows=4, cols=2)
            job_table.style = 'Table Grid'
            
            # Populate job info table
            job_data = [
                ('Job Title:', job.get('title', 'N/A')),
                ('Company:', job.get('company', 'N/A')),
                ('Location:', job.get('location', 'N/A')),
                ('Source:', job.get('source_type', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(job_data):
                job_table.cell(i, 0).text = label
                job_table.cell(i, 1).text = str(value)
            
            # Evaluation results section
            doc.add_heading('Evaluation Results', level=1)
            
            # Match score
            score_para = doc.add_paragraph()
            score_para.add_run('Match Score: ').bold = True
            score_para.add_run(f"{evaluation.get('match_score', 0)}/100")
            
            # Recommendation
            rec_para = doc.add_paragraph()
            rec_para.add_run('Recommendation: ').bold = True
            rec_para.add_run(evaluation.get('recommendation', 'N/A'))
            
            # Detailed feedback
            doc.add_heading('Detailed Analysis', level=1)
            feedback_text = evaluation.get('feedback', 'No detailed feedback available.')
            doc.add_paragraph(feedback_text)
            
            # Skills assessment if available
            if evaluation.get('skills_match'):
                doc.add_heading('Skills Assessment', level=2)
                skills_para = doc.add_paragraph()
                skills_para.add_run('Skills Match: ').bold = True
                skills_para.add_run(evaluation.get('skills_match', 'N/A'))
            
            # Experience assessment if available
            if evaluation.get('experience_match'):
                exp_para = doc.add_paragraph()
                exp_para.add_run('Experience Match: ').bold = True
                exp_para.add_run(evaluation.get('experience_match', 'N/A'))
            
            # Key strengths and concerns
            if evaluation.get('key_strengths'):
                doc.add_heading('Key Strengths', level=2)
                doc.add_paragraph(evaluation.get('key_strengths', ''))
            
            if evaluation.get('potential_concerns'):
                doc.add_heading('Potential Concerns', level=2)
                doc.add_paragraph(evaluation.get('potential_concerns', ''))
            
            # Footer with generation info
            doc.add_paragraph('\n')
            footer_para = doc.add_paragraph()
            footer_para.add_run('Generated on: ').italic = True
            footer_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M')).italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            self.logger.error(f"Error creating feedback document: {str(e)}")
            raise
    
    def create_resume_document(self, profile_data: Dict, job: Dict, 
                             customization_data: Dict) -> BytesIO:
        """Create a customized resume document in DOCX format"""
        try:
            doc = Document()
            
            # Header with name
            header = doc.add_heading('Professional Resume', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact information (placeholder)
            contact_para = doc.add_paragraph()
            contact_para.add_run('Email: candidate@email.com | Phone: (xxx) xxx-xxxx').italic = True
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Professional summary
            if customization_data.get('professional_summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(customization_data['professional_summary'])
            
            # Key skills
            if customization_data.get('key_skills'):
                doc.add_heading('Key Skills', level=1)
                skills_para = doc.add_paragraph()
                skills = customization_data['key_skills']
                if isinstance(skills, list):
                    skills_text = ' • '.join(skills)
                else:
                    skills_text = str(skills)
                skills_para.add_run(skills_text)
            
            # Experience section (placeholder)
            doc.add_heading('Professional Experience', level=1)
            exp_para = doc.add_paragraph()
            exp_para.add_run('[Experience details would be populated from profile data]').italic = True
            
            # Education section (placeholder)
            doc.add_heading('Education', level=1)
            edu_para = doc.add_paragraph()
            edu_para.add_run('[Education details would be populated from profile data]').italic = True
            
            # Customization note
            doc.add_paragraph('\n')
            custom_para = doc.add_paragraph()
            custom_para.add_run('Customized for: ').italic = True
            custom_para.add_run(f"{job.get('title', 'N/A')} at {job.get('company', 'N/A')}").italic = True
            custom_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            self.logger.error(f"Error creating resume document: {str(e)}")
            raise
    
    def create_summary_report(self, results: Dict) -> BytesIO:
        """Create a summary report of job evaluation results"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Job Evaluation Summary Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive summary
            doc.add_heading('Executive Summary', level=1)
            
            summary_stats = results.get('summary', {})
            
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Total jobs evaluated: {summary_stats.get('total_jobs', 0)}\n")
            stats_para.add_run(f"Average match score: {summary_stats.get('average_score', 0):.1f}\n")
            stats_para.add_run(f"High-scoring jobs (85+): {summary_stats.get('high_scoring_count', 0)}\n")
            
            # Top opportunities
            jobs = results.get('jobs', [])
            if jobs:
                doc.add_heading('Top Job Opportunities', level=1)
                
                # Create table for top jobs
                top_jobs = sorted(jobs, key=lambda x: x.get('match_score', 0), reverse=True)[:5]
                
                job_table = doc.add_table(rows=len(top_jobs) + 1, cols=4)
                job_table.style = 'Table Grid'
                
                # Headers
                headers = ['Score', 'Job Title', 'Company', 'Source']
                for i, header in enumerate(headers):
                    job_table.cell(0, i).text = header
                    job_table.cell(0, i).paragraphs[0].runs[0].bold = True
                
                # Job data
                for i, job in enumerate(top_jobs, 1):
                    job_table.cell(i, 0).text = str(job.get('match_score', 0))
                    job_table.cell(i, 1).text = job.get('title', 'N/A')
                    job_table.cell(i, 2).text = job.get('company', 'N/A')
                    job_table.cell(i, 3).text = job.get('source_type', 'N/A')
            
            # Generation info
            doc.add_paragraph('\n')
            footer_para = doc.add_paragraph()
            footer_para.add_run('Report generated on: ').italic = True
            footer_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M')).italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            self.logger.error(f"Error creating summary report: {str(e)}")
            raise
    
    def add_page_break(self, doc: Document):
        """Add a page break to the document"""
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE \\* MERGEFORMAT"
            run._r.append(instrText)
            
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
        except Exception as e:
            self.logger.error(f"Error adding page break: {str(e)}")